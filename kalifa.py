# with open('draft.N12.yosef.docx','r') as myf:

#     myf.read()
#     my
from docx import Document

document = Document('C:\\Users\\userx\\Desktop\\PROJECTS\\db_noflim\\src\\draftMaariv01.docx')
paragraphs = [p.text for p in document.paragraphs if p.text.strip() != '']
print(paragraphs[8])
# with open("forKalif3.docx",'w') as f:
#     f.write(str(paragraphs[::-1]))
#     f.close()
doc = Document()
paragraphs = paragraphs[::-1]
for i in paragraphs:

    doc.add_paragraph(i)
doc.save("kalif.docx")

# 'paragraphs' now contains all the paragraphs separated by "Enter"
